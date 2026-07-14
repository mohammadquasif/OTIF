"""Smoke test: DOCX upload -> scan -> approve -> rewrite -> DOCX/PDF export."""
import json
import sys
from io import BytesIO

sys.path.insert(0, "backend")

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.api.v1 import analysis, documents


def build_docx() -> BytesIO:
    doc = Document()
    doc.add_heading("Digital Transformation and SME Performance in India", level=1)
    sections = {
        "Abstract": "This study examines how digital transformation affects SME performance in India using survey and interview evidence.",
        "Introduction": "Digital transformation has become central to SME competitiveness. Smaller firms face financial, skill, and infrastructure constraints, and the literature often emphasizes large enterprises more than emerging-economy SMEs.",
        "Literature Review": "Vial (2019) defines digital transformation as a process where digital technologies trigger strategic organizational responses. Resource-based theory suggests that firms gain advantage when digital capabilities are valuable and difficult to imitate.",
        "Methodology": "The study uses a mixed-methods design with survey responses from SME owners and semi-structured interviews about adoption barriers.",
        "Results": "Firms with higher digital maturity reported stronger revenue growth and operational efficiency. Digital skill shortages remained a common barrier.",
        "Discussion": "The findings support the role of digital capability as a performance resource while showing that infrastructure and skills shape the gains.",
        "Conclusion": "Digital transformation can improve SME performance in India, but outcomes depend on skill development and implementation capability.",
        "References": "Vial, G. (2019). Understanding digital transformation: A review and a research agenda. Journal of Strategic Information Systems, 28(2), 118-144.",
    }
    for title, text in sections.items():
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def main() -> None:
    app = FastAPI()
    app.include_router(documents.router, prefix="/api/v1/documents")
    app.include_router(analysis.router, prefix="/api/v1/analysis")
    client = TestClient(app)

    result: dict = {}
    upload = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "existing_research.docx",
                build_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    result["upload_status"] = upload.status_code
    doc_id = upload.json().get("doc_id")
    result["doc_id"] = doc_id

    stages: list[str] = []
    final_scores = None
    plan: list[dict] = []
    with client.stream(
        "POST",
        f"/api/v1/analysis/run/{doc_id}",
        json={"doc_type": "research_paper", "norm": "apa7", "pace": "fast"},
        timeout=300,
    ) as response:
        result["scan_status"] = response.status_code
        for line in response.iter_lines():
            if not line or not line.startswith("data: "):
                continue
            payload = json.loads(line[6:])
            stages.append(payload.get("stage", ""))
            if payload.get("stage") == "scores_ready":
                final_scores = payload.get("scores")
                plan = payload.get("improvement_plan") or []

    result["scan_event_count"] = len(stages)
    result["scan_unique_stages"] = len(set(stages))
    result["scan_complete"] = "complete" in stages and "error" not in stages
    result["scores"] = final_scores
    result["plan_count"] = len(plan)
    result["plan_titles"] = [item.get("title") for item in plan[:5]]

    approved_ids = [item["id"] for item in plan[:3]]
    approve = client.post(
        "/api/v1/analysis/approve-rewrite",
        json={
            "doc_id": doc_id,
            "approved_item_ids": approved_ids,
            "doc_type": "research_paper",
            "norm": "apa7",
            "design_theme": "mono_formal",
            "output_formats": ["docx", "pdf"],
        },
    )
    result["approve_status"] = approve.status_code
    approval_json = approve.json()
    result["rewrite_status"] = approval_json.get("rewrite_status")
    result["preview_chars"] = len(approval_json.get("rewrite_preview") or "")

    chapter = client.get(
        f"/api/v1/analysis/chapter-editor/{doc_id}",
        params={"doc_type": "research_paper", "norm": "apa7"},
    )
    result["chapter_status"] = chapter.status_code
    chapters = chapter.json().get("chapters", [])
    result["chapter_count"] = len(chapters)
    chapter_payload = [
        {
            "id": c["id"],
            "title": c["title"],
            "original_text": c["original_text"],
            "edited_text": c["edited_text"],
        }
        for c in chapters
    ]

    rewrite = client.post(
        "/api/v1/analysis/rewrite-full-document",
        json={
            "doc_id": doc_id,
            "chapters": chapter_payload,
            "approved_item_ids": approved_ids,
            "doc_type": "research_paper",
            "norm": "apa7",
            "design_theme": "mono_formal",
        },
        timeout=300,
    )
    result["full_rewrite_status"] = rewrite.status_code
    rewrite_json = rewrite.json()
    result["full_rewrite_chapters"] = len(rewrite_json.get("chapters", []))
    result["rewrite_words_original"] = rewrite_json.get("total_words_original")
    result["rewrite_words_rewritten"] = rewrite_json.get("total_words_rewritten")
    result["rewrite_warnings"] = rewrite_json.get("warnings")
    result["rewrite_error"] = rewrite_json.get("detail")
    export_chapters = [
        {
            "id": c["id"],
            "title": c["title"],
            "original_text": c["original_text"],
            "edited_text": c["rewritten_text"],
        }
        for c in rewrite_json.get("chapters", [])
    ]

    finalize = client.post(
        "/api/v1/analysis/finalize-thesis",
        json={
            "doc_id": doc_id,
            "chapters": export_chapters,
            "doc_type": "research_paper",
            "norm": "apa7",
            "design_theme": "mono_formal",
            "output_formats": ["docx", "pdf"],
        },
        timeout=180,
    )
    result["finalize_status"] = finalize.status_code
    final_json = finalize.json()
    result["artifact_summary"] = [
        {
            "format": artifact.get("format"),
            "filename": artifact.get("filename"),
            "size_bytes": artifact.get("size_bytes"),
        }
        for artifact in final_json.get("artifacts", [])
    ]
    result["preservation_report"] = final_json.get("preservation_report")
    result["field_update_status"] = final_json.get("field_update_status")

    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
