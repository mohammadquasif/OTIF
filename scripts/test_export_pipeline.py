"""Smoke test for OTIF final DOCX/PDF export.

Run from the repository root with the backend environment active:
    python scripts/test_export_pipeline.py

This verifies the production export dependencies and creates a tiny local
DOCX/PDF/certificate package without contacting any network service.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
BACKEND = ROOT / "backend"
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(BACKEND))

from app.export.thesis_exporter import (  # noqa: E402
    build_integrity_certificate,
    create_certificate_markdown,
    create_docx,
    create_preserved_docx,
    create_pdf,
    normalize_hex_color,
    render_diagram_image,
    resolve_theme,
)


def main() -> None:
    normalize_hex_color("#1F4E79")
    theme = resolve_theme("classic_blue", "#1F4E79")
    chapters = [
        {
            "id": "chapter-1",
            "title": "Chapter 1: Introduction",
            "original_text": "This study examines a local-first academic integrity workflow.",
            "edited_text": (
                "This study examines a local-first academic integrity workflow. "
                "The contribution is positioned as a defensible verification and revision process."
            ),
        }
    ]
    approval = {
        "approved_item_ids": ["academic-voice"],
        "approved_items": [
            {
                "id": "academic-voice",
                "title": "Improve academic voice",
                "action": "Strengthen scholarly clarity while preserving claims.",
            }
        ],
        "active_provider": "ollama",
        "active_model": "llama3.3:latest",
    }
    certificate = build_integrity_certificate(
        doc_type="thesis",
        norm="ugc",
        approval=approval,
        before_scores={"overall_preflight": 61.5, "citation_quality": 50},
        after_scores={"overall_preflight": 74.2, "citation_quality": 67},
        chapter_count=len(chapters),
    )

    with tempfile.TemporaryDirectory(prefix="otif_export_smoke_") as tmp:
        out_dir = Path(tmp)
        from docx import Document

        source_doc = Document()
        source_doc.add_heading("Chapter 1: Introduction", level=1)
        source_doc.add_paragraph("This study examines a local-first academic integrity workflow.")
        table = source_doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Status"
        table.cell(1, 0).text = "Preserved table"
        table.cell(1, 1).text = "Expected"
        source_docx = out_dir / "otif_source_roundtrip.docx"
        source_doc.save(source_docx)

        diagram = render_diagram_image(
            diagram_source="flowchart TD\n  A[Upload] --> B[Verify]\n  B --> C[Finalize]",
            diagram_caption="Figure 1. OTIF verification flow",
            output_dir=out_dir,
            accent=theme["accent"],
            stem="smoke",
        )
        docx = create_docx(
            output_path=out_dir / "otif_export_smoke.docx",
            document_title="OTIF Export Smoke Test",
            chapters=chapters,
            doc_type="thesis",
            norm="ugc",
            design_theme="classic_blue",
            design_accent_hex="#1F4E79",
            approval=approval,
            certificate=certificate,
            diagram_source="flowchart TD\n  A[Upload] --> B[Verify]\n  B --> C[Finalize]",
            diagram_caption="Figure 1. OTIF verification flow",
            diagram_image_path=diagram,
        )
        pdf = create_pdf(
            output_path=out_dir / "otif_export_smoke.pdf",
            document_title="OTIF Export Smoke Test",
            chapters=chapters,
            doc_type="thesis",
            norm="ugc",
            design_theme="classic_blue",
            design_accent_hex="#1F4E79",
            approval=approval,
            certificate=certificate,
            diagram_source="flowchart TD\n  A[Upload] --> B[Verify]\n  B --> C[Finalize]",
            diagram_caption="Figure 1. OTIF verification flow",
            diagram_image_path=diagram,
        )
        cert = create_certificate_markdown(
            output_path=out_dir / "otif_export_certificate.md",
            document_title="OTIF Export Smoke Test",
            certificate=certificate,
            approval=approval,
        )
        preserved, preservation_report = create_preserved_docx(
            source_path=source_docx,
            output_path=out_dir / "otif_preserved_roundtrip.docx",
            chapters=chapters,
            certificate=certificate,
            approval=approval,
            diagram_source=None,
            diagram_caption=None,
            diagram_image_path=None,
        )
        for artifact in [docx, pdf, cert]:
            if artifact.size_bytes <= 0 or not artifact.path.exists():
                raise AssertionError(f"Artifact failed: {artifact.path}")
            print(f"OK {artifact.format}: {artifact.path.name} ({artifact.size_bytes} bytes)")
        if preserved.size_bytes <= 0 or preservation_report["tables_preserved"] < 1:
            raise AssertionError("DOCX round-trip preservation failed")
        print(
            "OK preserved docx: "
            f"{preserved.path.name} ({preserved.size_bytes} bytes), "
            f"chapters matched={preservation_report['chapters_matched']}, "
            f"tables preserved={preservation_report['tables_preserved']}"
        )
        if diagram:
            print(f"OK diagram image: {diagram.name} ({diagram.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
