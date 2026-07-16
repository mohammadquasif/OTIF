"""Final thesis package generation.

The exporter builds user-approved, text-faithful DOCX/PDF artifacts from the
chapter editor state. It keeps imports lazy so the desktop backend can start
even when an optional packaging dependency needs repair.
"""
from __future__ import annotations

import re
import shutil
import subprocess
from html import unescape
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape


@dataclass(frozen=True)
class ExportArtifact:
    format: str
    filename: str
    path: Path
    size_bytes: int


THEMES: dict[str, dict[str, str]] = {
    "classic_blue": {"accent": "1f4e79", "muted": "eef4fb", "name": "Classic Blue"},
    "emerald_research": {"accent": "0f766e", "muted": "ecfdf5", "name": "Emerald Research"},
    "emerald_academic": {"accent": "0f766e", "muted": "ecfdf5", "name": "Emerald Academic"},
    "monochrome": {"accent": "111827", "muted": "f3f4f6", "name": "Monochrome"},
    "mono_formal": {"accent": "111827", "muted": "f3f4f6", "name": "Mono Formal"},
    "royal_purple": {"accent": "5b21b6", "muted": "f5f3ff", "name": "Royal Purple"},
    "maroon_submission": {"accent": "7f1d1d", "muted": "fef2f2", "name": "Maroon Submission"},
}

HEX_COLOR_RE = re.compile(r"^#?[0-9A-Fa-f]{6}$")


NORM_LABELS = {
    "apa7": "APA 7",
    "ugc": "UGC thesis format",
    "ieee": "IEEE",
    "harvard": "Harvard",
    "springer": "Springer",
    "elsevier": "Elsevier",
    "european_thesis": "European thesis format",
}


def export_dir(base_uploads_dir: Path, doc_id: str) -> Path:
    path = base_uploads_dir / f"{doc_id}_exports"
    path.mkdir(parents=True, exist_ok=True)
    return path


# ── Real ToC / LOT / LOF generation ─────────────────────────────

WORDS_PER_PAGE = {"ugc": 320, "apa7": 350, "ieee": 450, "harvard": 350, "springer": 400, "elsevier": 400, "default": 380}

def _extract_toc_entries(chapters: list[dict], norm: str) -> list[dict]:
    """Parse chapter content and extract headings, figures, and tables with estimated page numbers."""
    wpp = WORDS_PER_PAGE.get(norm, WORDS_PER_PAGE["default"])
    entries: list[dict] = []
    figures: list[dict] = []
    tables: list[dict] = []

    cumulative_pages = 1  # Start after title page
    # Front matter pages
    cumulative_pages += 2  # title page + subtitle

    for ch_idx, chapter in enumerate(chapters):
        title = str(chapter.get("title") or f"Chapter {ch_idx + 1}")
        text = str(chapter.get("edited_text") or chapter.get("text") or "")
        word_count = len(text.split())

        # Heading entry for ToC
        entries.append({
            "level": 1, "title": title, "page": cumulative_pages,
            "words": word_count, "chapter_index": ch_idx,
        })

        # Extract sub-headings and figures/tables from HTML or markdown
        # Look for markdown headings ## and ###
        for m in re.finditer(r"^(#{2,4})\s+(.+)$", text, re.MULTILINE):
            level = len(m.group(1))
            sub_title = m.group(2).strip()
            sub_words = len(text[m.start():].split())
            sub_page = cumulative_pages + max(0, (word_count - sub_words) // wpp)
            entries.append({
                "level": min(level, 3), "title": sub_title, "page": max(cumulative_pages, sub_page),
                "words": 0, "chapter_index": ch_idx,
            })

        # Extract diagrams/figures
        for m in re.finditer(r"(?:```mermaid\n.*?```|Figure\s+\d+[.:]\s*(.+?)(?:\n|$))", text, re.IGNORECASE | re.DOTALL):
            caption = m.group(1).strip() if m.lastindex else "Diagram"
            fig_pos = len(text[:m.start()].split())
            fig_page = cumulative_pages + max(0, fig_pos // wpp)
            figures.append({
                "caption": caption[:120], "page": max(cumulative_pages, fig_page),
                "chapter": title,
            })

        # Extract tables
        for m in re.finditer(r"Table\s+\d+[.:]\s*(.+?)(?:\n|$)", text, re.IGNORECASE):
            caption = m.group(1).strip()
            tbl_pos = len(text[:m.start()].split())
            tbl_page = cumulative_pages + max(0, tbl_pos // wpp)
            tables.append({
                "caption": caption[:120], "page": max(cumulative_pages, tbl_page),
                "chapter": title,
            })

        # Advance page counter
        chapter_pages = max(1, word_count // wpp)
        cumulative_pages += chapter_pages

    return entries, figures, tables


def build_toc_text(entries: list[dict]) -> str:
    """Build a formatted Table of Contents string with dot leaders and page numbers."""
    lines = ["Table of Contents", "=" * 50, ""]
    for entry in entries:
        indent = "  " * (entry["level"] - 1)
        title = entry["title"][:80]
        page = str(entry["page"])
        dots = "." * max(2, 55 - len(indent) - len(title) - len(page))
        lines.append(f"{indent}{title} {dots} {page}")
    return "\n".join(lines)


def build_lot_text(figures: list[dict]) -> str:
    """Build a formatted List of Figures with captions and page numbers."""
    lines = ["List of Figures", "=" * 50, ""]
    for i, fig in enumerate(figures, 1):
        caption = fig["caption"][:90]
        page = str(fig["page"])
        dots = "." * max(2, 55 - len(f"Figure {i}: ") - len(caption) - len(page))
        lines.append(f"Figure {i}: {caption} {dots} {page}")
    if not figures:
        lines.append("(No figures in this document)")
    return "\n".join(lines)


def build_lol_text(tables: list[dict]) -> str:
    """Build a formatted List of Tables with captions and page numbers."""
    lines = ["List of Tables", "=" * 50, ""]
    for i, tbl in enumerate(tables, 1):
        caption = tbl["caption"][:90]
        page = str(tbl["page"])
        dots = "." * max(2, 55 - len(f"Table {i}: ") - len(caption) - len(page))
        lines.append(f"Table {i}: {caption} {dots} {page}")
    if not tables:
        lines.append("(No tables in this document)")
    return "\n".join(lines)


def safe_filename(value: str, fallback: str = "otif_thesis") -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._-")
    return stem[:120] or fallback


def normalize_hex_color(value: str | None) -> str | None:
    if not value:
        return None
    candidate = value.strip()
    if not HEX_COLOR_RE.fullmatch(candidate):
        raise ValueError("Custom design color must be a 6-digit hex value such as #1F4E79.")
    return candidate.lstrip("#").lower()


def _muted_from_accent(accent: str) -> str:
    red = int(accent[0:2], 16)
    green = int(accent[2:4], 16)
    blue = int(accent[4:6], 16)
    blend = lambda channel: int(channel * 0.08 + 255 * 0.92)
    return f"{blend(red):02x}{blend(green):02x}{blend(blue):02x}"


def resolve_theme(design_theme: str, design_accent_hex: str | None = None) -> dict[str, str]:
    theme = dict(THEMES.get(design_theme, THEMES["classic_blue"]))
    accent = normalize_hex_color(design_accent_hex)
    if accent:
        theme["accent"] = accent
        theme["muted"] = _muted_from_accent(accent)
        theme["name"] = f"Custom #{accent.upper()}"
    return theme


def _mark_update_fields_on_open(doc: Any) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")
    except Exception:
        pass


def _add_field_paragraph(doc: Any, instr: str, fallback: str) -> Any:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instr)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = fallback
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    return paragraph


def _add_caption(doc: Any, label: str, caption: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    try:
        paragraph.style = doc.styles["Caption"]
    except KeyError:
        pass
    paragraph.add_run(f"{label} ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f"SEQ {label} \\* ARABIC")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    paragraph.add_run(f": {caption.strip()}")


def update_docx_fields_with_word(docx_path: Path) -> bool:
    """Update TOC/TOF/LOT/page fields in-place when Microsoft Word is available."""
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False

    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        try:
            doc.Repaginate()
        except Exception:
            pass
        for toc in doc.TablesOfContents:
            toc.Update()
        for table in doc.TablesOfFigures:
            table.Update()
        try:
            doc.Fields.Update()
        except Exception:
            pass
        doc.Save()
        return True
    except Exception:
        return False
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def compile_chapter_text(chapters: list[dict[str, Any]]) -> str:
    chunks: list[str] = []
    for chapter in chapters:
        title = str(chapter.get("title") or "Untitled chapter").strip()
        text = rich_text_to_plain_text(str(chapter.get("edited_text") or chapter.get("text") or ""))
        if title:
            chunks.append(title)
        if text:
            chunks.append(text)
    return "\n\n".join(chunks).strip()


def rich_text_to_plain_text(value: str) -> str:
    """Convert TipTap/editor HTML or plain text into export-ready manuscript text."""
    text = str(value or "").strip()
    if not text:
        return ""
    if "<" not in text or ">" not in text:
        return text

    text = re.sub(r"(?i)<\s*br\s*/?\s*>", "\n", text)
    text = re.sub(r"(?i)</\s*(p|div|h[1-6]|li|blockquote|tr)\s*>", "\n\n", text)
    text = re.sub(r"(?i)<\s*li[^>]*>", "- ", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _heading_level(paragraph: Any) -> int | None:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    match = re.match(r"Heading\s+(\d+)", style_name, flags=re.I)
    if match:
        return int(match.group(1))
    try:
        outline = paragraph._p.pPr.outlineLvl
        if outline is not None and outline.val is not None:
            return int(outline.val) + 1
    except Exception:
        pass
    return None


def _is_chapter_heading(paragraph: Any) -> bool:
    text = paragraph.text.strip()
    level = _heading_level(paragraph)
    if level == 1:
        return True
    return bool(
        re.match(
            r"(?i)^(chapter\s+\d+|chapter\s+[ivxlcdm]+|abstract|introduction|literature review|methodology|methods|results|discussion|conclusion|references)\b",
            text,
        )
    )


def _matches_chapter_title(paragraph_text: str, chapter_title: str) -> bool:
    left = _normalize_text(paragraph_text)
    right = _normalize_text(chapter_title)
    if not left or not right:
        return False
    return left == right or left in right or right in left


def _has_complex_content(paragraph: Any) -> bool:
    element = paragraph._p
    return bool(
        element.xpath(".//w:drawing")
        or element.xpath(".//w:pict")
        or element.xpath(".//w:object")
        or element.xpath(".//w:fldChar")
        or element.xpath(".//w:instrText")
        or element.xpath(".//w:hyperlink")
    )


def _is_protected_paragraph(paragraph: Any) -> bool:
    text = paragraph.text.strip()
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    if not text:
        return True
    if _is_chapter_heading(paragraph):
        return True
    if _has_complex_content(paragraph):
        return True
    if re.match(r"(?i)^(figure|fig\.|table|appendix)\s+[\w\d]+[:.\-\s]", text):
        return True
    if "caption" in style_name.lower():
        return True
    if re.match(r"(?i)^(references|bibliography|works cited)\s*$", text):
        return True
    return False


def _clear_and_set_paragraph_text(paragraph: Any, text: str) -> None:
    p = paragraph._p
    p_pr = p.pPr
    for child in list(p):
        if child is not p_pr:
            p.remove(child)
    paragraph.add_run(text)


def _delete_paragraph(paragraph: Any) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _split_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n\s*\n+", text.strip()) if block.strip()]


def _append_otif_back_matter(
    *,
    doc: Any,
    certificate: dict[str, Any],
    approval: dict[str, Any] | None,
    diagram_source: str | None,
    diagram_caption: str | None,
    diagram_image_path: Path | None,
    preservation_report: dict[str, Any],
) -> None:
    from docx.shared import Inches

    doc.add_page_break()
    doc.add_heading("OTIF Final Integrity Certificate", level=1)
    doc.add_paragraph(certificate["statement"])
    doc.add_paragraph(f"Generated: {certificate['generated_at']}")
    doc.add_paragraph(f"Target format: {certificate['target_format']}")
    doc.add_paragraph(f"Chapters finalized: {certificate['chapter_count']}")
    doc.add_paragraph(
        "Preservation mode: original DOCX structure retained; editable chapter body paragraphs revised in place."
    )

    doc.add_heading("Preservation Summary", level=2)
    for key, value in preservation_report.items():
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    approved_items = approval.get("approved_items", []) if approval else []
    if approved_items:
        doc.add_heading("Approved Improvement Scope", level=2)
        for item in approved_items:
            doc.add_paragraph(f"{item.get('title', 'Improvement item')}: {item.get('action', '')}")

    if diagram_source:
        doc.add_page_break()
        doc.add_heading("Approved Diagram", level=1)
        if diagram_caption:
            doc.add_paragraph(diagram_caption)
        if diagram_image_path and diagram_image_path.exists():
            doc.add_picture(str(diagram_image_path), width=Inches(6.4))
        else:
            code = doc.add_paragraph()
            code_run = code.add_run(diagram_source)
            code_run.font.name = "Courier New"


def create_preserved_docx(
    *,
    source_path: Path,
    output_path: Path,
    chapters: list[dict[str, Any]],
    certificate: dict[str, Any],
    approval: dict[str, Any] | None,
    diagram_source: str | None = None,
    diagram_caption: str | None = None,
    diagram_image_path: Path | None = None,
) -> tuple[ExportArtifact, dict[str, Any]]:
    """Copy an original DOCX and apply chapter edits while preserving structure.

    The method edits only normal body paragraphs inside matched chapter ranges.
    Tables, images, captions, hyperlinks, TOC fields, and front matter are not
    deleted or regenerated.
    """
    from docx import Document

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_path, output_path)
    doc = Document(output_path)
    paragraphs = list(doc.paragraphs)
    chapter_heading_indexes = [idx for idx, paragraph in enumerate(paragraphs) if _is_chapter_heading(paragraph)]
    matched = 0
    inserted = 0
    replaced = 0
    removed = 0
    protected_seen = 0
    missing_titles: list[str] = []

    for chapter in chapters:
        title = str(chapter.get("title") or "").strip()
        edited_blocks = _split_blocks(rich_text_to_plain_text(str(chapter.get("edited_text") or chapter.get("text") or "")))
        if not title or not edited_blocks:
            continue

        start_index = next(
            (
                idx
                for idx in chapter_heading_indexes
                if _matches_chapter_title(paragraphs[idx].text, title)
            ),
            None,
        )
        if start_index is None:
            missing_titles.append(title)
            continue

        next_heading = next((idx for idx in chapter_heading_indexes if idx > start_index), len(paragraphs))
        range_paragraphs = paragraphs[start_index + 1 : next_heading]
        editable = [paragraph for paragraph in range_paragraphs if not _is_protected_paragraph(paragraph)]
        protected_seen += len(range_paragraphs) - len(editable)
        matched += 1

        for paragraph, block in zip(editable, edited_blocks):
            _clear_and_set_paragraph_text(paragraph, block)
            replaced += 1

        if len(edited_blocks) > len(editable):
            anchor = paragraphs[next_heading] if next_heading < len(paragraphs) else None
            for block in edited_blocks[len(editable) :]:
                if anchor is not None:
                    anchor.insert_paragraph_before(block)
                else:
                    doc.add_paragraph(block)
                inserted += 1

        for paragraph in editable[len(edited_blocks) :]:
            _delete_paragraph(paragraph)
            removed += 1

    report = {
        "mode": "docx_round_trip_preservation",
        "chapters_requested": len(chapters),
        "chapters_matched": matched,
        "missing_chapter_titles": missing_titles,
        "body_paragraphs_replaced": replaced,
        "body_paragraphs_inserted": inserted,
        "body_paragraphs_removed": removed,
        "protected_paragraphs_preserved": protected_seen,
        "tables_preserved": len(doc.tables),
        "paragraphs_preserved_or_updated": len(doc.paragraphs),
    }

    _append_otif_back_matter(
        doc=doc,
        certificate=certificate,
        approval=approval,
        diagram_source=diagram_source,
        diagram_caption=diagram_caption,
        diagram_image_path=diagram_image_path,
        preservation_report=report,
    )
    doc.core_properties.comments = "Generated locally by OTIF with DOCX round-trip preservation."
    _mark_update_fields_on_open(doc)
    doc.save(output_path)
    return ExportArtifact("docx", output_path.name, output_path, output_path.stat().st_size), report


def build_integrity_certificate(
    *,
    doc_type: str,
    norm: str,
    approval: dict[str, Any] | None,
    before_scores: dict[str, Any],
    after_scores: dict[str, Any],
    chapter_count: int,
) -> dict[str, Any]:
    return {
        "title": "OTIF Final Integrity Certificate",
        "generated_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "doc_type": doc_type,
        "target_format": NORM_LABELS.get(norm, norm.upper()),
        "chapter_count": chapter_count,
        "approved_items": approval.get("approved_item_ids", []) if approval else [],
        "active_provider": approval.get("active_provider") if approval else None,
        "active_model": approval.get("active_model") if approval else None,
        "before": before_scores,
        "after": after_scores,
        "statement": (
            "This package was generated from user-reviewed chapter edits and an approved OTIF "
            "improvement plan. Citations and source claims remain the author's responsibility."
        ),
    }


def create_docx(
    *,
    output_path: Path,
    document_title: str,
    chapters: list[dict[str, Any]],
    doc_type: str,
    norm: str,
    design_theme: str,
    design_accent_hex: str | None,
    approval: dict[str, Any] | None,
    certificate: dict[str, Any],
    diagram_source: str | None = None,
    diagram_caption: str | None = None,
    diagram_image_path: Path | None = None,
) -> ExportArtifact:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    theme = resolve_theme(design_theme, design_accent_hex)
    accent = RGBColor.from_string(theme["accent"])
    muted = RGBColor.from_string(theme.get("muted", "888888"))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Page border (accent color, subtle) ──────────────────
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sect_pr = section._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for edge_name in ('top', 'left', 'bottom', 'right'):
        edge = OxmlElement(f'w:{edge_name}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '4')
        edge.set(qn('w:space'), '18')
        edge.set(qn('w:color'), theme['accent'])
        pg_borders.append(edge)
    sect_pr.append(pg_borders)

    # ── Footer with page numbers ────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp_run = fp.add_run()
    fp_run.font.size = Pt(10)
    fp_run.font.color.rgb = muted
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    fp_run._r.append(fld_char_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fp_run._r.append(instr)
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    fp_run._r.append(fld_char_separate)
    page_text = OxmlElement('w:t')
    page_text.text = '1'
    fp_run._r.append(page_text)
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    fp_run._r.append(fld_char_end)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman" if norm in {"ugc", "apa7"} else "Arial"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5 if norm == "ugc" else 1.15
    normal.paragraph_format.space_after = Pt(8)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = accent
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(document_title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = accent

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{doc_type.title()} | {NORM_LABELS.get(norm, norm.upper())} | {theme['name']}")

    doc.add_paragraph()
    # ── Real ToC / LOT / LOF with computed page numbers ──
    toc_entries, figures_list, tables_list = _extract_toc_entries(chapters, norm)

    doc.add_heading("Table of Contents", level=1)
    for entry in toc_entries:
        indent = "    " * (entry["level"] - 1)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(1)
        # Use proper tab-stop dot leader for cleaner TOC
        p_pr_toc = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9360')
        tabs.append(tab)
        p_pr_toc.append(tabs)
        text_run = p.add_run(f"{indent}{entry['title']}\t{entry['page']}")
        text_run.font.size = Pt(11)

    doc.add_page_break()
    doc.add_heading("List of Tables", level=1)
    if tables_list:
        for i, tbl in enumerate(tables_list, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"Table {i}: {tbl['caption'][:100]}").font.size = Pt(10)
            dots = p.add_run(f"  . . . . . . . . . . . . . . . . . . . .  {tbl['page']}")
            dots.font.size = Pt(9)
            dots.font.color.rgb = muted
    else:
        doc.add_paragraph("(No tables in this document)")

    doc.add_page_break()
    doc.add_heading("List of Figures", level=1)
    if figures_list:
        for i, fig in enumerate(figures_list, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"Figure {i}: {fig['caption'][:100]}").font.size = Pt(10)
            dots = p.add_run(f"  . . . . . . . . . . . . . . . . . . . .  {fig['page']}")
            dots.font.size = Pt(9)
            dots.font.color.rgb = muted
    else:
        doc.add_paragraph("(No figures in this document)")

    doc.add_page_break()

    for chapter in chapters:
        heading = doc.add_heading(str(chapter.get("title") or "Untitled chapter"), level=1)
        # Add accent-colored bottom border to heading
        p_pr = heading._p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), theme['accent'])
        bottom.set(qn('w:space'), '4')
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

        body = rich_text_to_plain_text(str(chapter.get("edited_text") or chapter.get("text") or ""))
        for block in re.split(r"\n\s*\n+", body):
            block = block.strip()
            if not block:
                continue
            caption_match = re.match(r"^(table|figure)\s+\d*[:.\-\s]*(.*)", block, flags=re.I)
            if caption_match:
                _add_caption(doc, caption_match.group(1).title(), caption_match.group(2) or block)
                continue
            # Detect sub-headings (## / ### markdown)
            h2_match = re.match(r'^##\s+(.+)$', block)
            if h2_match:
                doc.add_heading(h2_match.group(1), level=2)
                continue
            h3_match = re.match(r'^###\s+(.+)$', block)
            if h3_match:
                doc.add_heading(h3_match.group(1), level=3)
                continue
            doc.add_paragraph(block)

    if diagram_source:
        doc.add_page_break()
        doc.add_heading("Approved Diagram", level=1)
        if diagram_caption:
            _add_caption(doc, "Figure", diagram_caption)
        if diagram_image_path and diagram_image_path.exists():
            doc.add_picture(str(diagram_image_path), width=Inches(6.4))
        else:
            code = doc.add_paragraph()
            code_run = code.add_run(diagram_source)
            code_run.font.name = "Courier New"
            code_run.font.size = Pt(9)

    doc.add_page_break()
    doc.add_heading("Approved Improvement Plan", level=1)
    approved_items = approval.get("approved_items", []) if approval else []
    if approved_items:
        for item in approved_items:
            doc.add_heading(str(item.get("title", "Improvement item")), level=2)
            doc.add_paragraph(str(item.get("action", "")))
            evidence = str(item.get("evidence", "")).strip()
            if evidence:
                doc.add_paragraph(f"Evidence: {evidence}")
    else:
        doc.add_paragraph("No approved improvement items were recorded for this export.")

    doc.add_page_break()
    doc.add_heading(certificate["title"], level=1)
    doc.add_paragraph(certificate["statement"])
    doc.add_paragraph(f"Generated: {certificate['generated_at']}")
    doc.add_paragraph(f"Target format: {certificate['target_format']}")
    doc.add_paragraph(f"Chapters finalized: {certificate['chapter_count']}")
    doc.add_heading("Before and After Scores", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Metric"
    header[1].text = "Before"
    header[2].text = "After"
    keys = sorted(set(certificate["before"].keys()) | set(certificate["after"].keys()))
    for key in keys:
        row = table.add_row().cells
        row[0].text = key.replace("_", " ").title()
        row[1].text = str(certificate["before"].get(key, ""))
        row[2].text = str(certificate["after"].get(key, ""))

    doc.core_properties.title = document_title
    doc.core_properties.subject = "OTIF finalized academic document"
    doc.core_properties.comments = "Generated locally by OTIF after user approval."
    _mark_update_fields_on_open(doc)
    doc.save(output_path)
    return ExportArtifact("docx", output_path.name, output_path, output_path.stat().st_size)


def _find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> ExportArtifact | None:
    """Try to render a DOCX to PDF with an office engine for visual parity."""
    update_docx_fields_with_word(docx_path)
    soffice = _find_soffice()
    if soffice:
        tmp_dir = pdf_path.parent / "_pdf_convert"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_dir),
                    str(docx_path),
                ],
                check=True,
                timeout=90,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            converted = tmp_dir / f"{docx_path.stem}.pdf"
            if converted.exists():
                if pdf_path.exists():
                    pdf_path.unlink()
                converted.replace(pdf_path)
                return ExportArtifact("pdf", pdf_path.name, pdf_path, pdf_path.stat().st_size)
        except Exception:
            pass
        finally:
            try:
                if tmp_dir.exists():
                    shutil.rmtree(tmp_dir)
            except OSError:
                pass

    try:
        import win32com.client  # type: ignore
    except Exception:
        return None

    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
        if pdf_path.exists():
            return ExportArtifact("pdf", pdf_path.name, pdf_path, pdf_path.stat().st_size)
    except Exception:
        return None
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
    return None


def _render_diagram_card_svg(diagram_source: str, caption: str | None, accent: str) -> bytes:
    lines = [line[:110] for line in diagram_source.strip().splitlines()[:28]]
    width = 1200
    line_height = 28
    caption_height = 50 if caption else 0
    height = max(360, 120 + caption_height + len(lines) * line_height)
    text_lines = "\n".join(
        f'<text x="64" y="{150 + caption_height + idx * line_height}" '
        f'font-family="Consolas, Courier New, monospace" font-size="18" fill="#111827">'
        f"{xml_escape(line)}</text>"
        for idx, line in enumerate(lines)
    )
    caption_markup = (
        f'<text x="64" y="112" font-family="Arial, sans-serif" font-size="22" '
        f'font-weight="700" fill="#111827">{xml_escape(caption or "")}</text>'
        if caption
        else ""
    )
    return f"""<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">
  <rect width="100%" height="100%" fill="#ffffff"/>
  <rect x="28" y="28" width="{width - 56}" height="{height - 56}" rx="18" fill="#f9fafb" stroke="#{accent}" stroke-width="4"/>
  <text x="64" y="72" font-family="Arial, sans-serif" font-size="28" font-weight="700" fill="#{accent}">Approved OTIF Diagram</text>
  {caption_markup}
  {text_lines}
</svg>""".encode("utf-8")


def render_diagram_image(
    *,
    diagram_source: str | None,
    diagram_caption: str | None,
    output_dir: Path,
    accent: str,
    stem: str,
) -> Path | None:
    if not diagram_source or not diagram_source.strip():
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    mmdc = shutil.which("mmdc")
    source_path = output_dir / f"{stem}_diagram.mmd"
    png_path = output_dir / f"{stem}_diagram.png"
    source_path.write_text(diagram_source, encoding="utf-8")
    if mmdc:
        try:
            subprocess.run(
                [mmdc, "-i", str(source_path), "-o", str(png_path), "-b", "transparent"],
                check=True,
                timeout=90,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if png_path.exists():
                return png_path
        except Exception:
            pass

    try:
        import fitz

        svg = _render_diagram_card_svg(diagram_source, diagram_caption, accent)
        svg_doc = fitz.open(stream=svg, filetype="svg")
        page = svg_doc[0]
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        pixmap.save(png_path)
        svg_doc.close()
        if png_path.exists():
            return png_path
    except Exception:
        return None
    return None


def create_pdf(
    *,
    output_path: Path,
    document_title: str,
    chapters: list[dict[str, Any]],
    doc_type: str,
    norm: str,
    design_theme: str,
    design_accent_hex: str | None,
    approval: dict[str, Any] | None,
    certificate: dict[str, Any],
    diagram_source: str | None = None,
    diagram_caption: str | None = None,
    diagram_image_path: Path | None = None,
) -> ExportArtifact:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from xml.sax.saxutils import escape

    theme = resolve_theme(design_theme, design_accent_hex)
    accent = colors.HexColor(f"#{theme['accent']}")
    styles = getSampleStyleSheet()
    body_font = "Times-Roman" if norm in {"ugc", "apa7"} else "Helvetica"
    body = ParagraphStyle(
        "OTIFBody",
        parent=styles["BodyText"],
        fontName=body_font,
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )
    h1 = ParagraphStyle(
        "OTIFHeading1",
        parent=styles["Heading1"],
        textColor=accent,
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        spaceBefore=12,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "OTIFHeading2",
        parent=styles["Heading2"],
        textColor=accent,
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
    )
    code_style = ParagraphStyle(
        "OTIFCode",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=8,
        leading=11,
        backColor=colors.HexColor(f"#{theme['muted']}"),
        borderColor=accent,
        borderWidth=0.5,
        borderPadding=6,
    )

    story: list[Any] = []
    story.append(Paragraph(escape(document_title), h1))
    story.append(Paragraph(escape(f"{doc_type.title()} | {NORM_LABELS.get(norm, norm.upper())} | {theme['name']}"), body))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph("Table of Contents", h1))
    for idx, chapter in enumerate(chapters, start=1):
        story.append(Paragraph(escape(f"{idx}. {chapter.get('title') or 'Untitled chapter'}"), body))
    story.append(PageBreak())

    for chapter in chapters:
        story.append(Paragraph(escape(str(chapter.get("title") or "Untitled chapter")), h1))
        body_text = rich_text_to_plain_text(str(chapter.get("edited_text") or chapter.get("text") or ""))
        for block in re.split(r"\n\s*\n+", body_text):
            block = block.strip()
            if block:
                story.append(Paragraph(escape(block).replace("\n", "<br/>"), body))

    if diagram_source:
        story.append(PageBreak())
        story.append(Paragraph("Approved Diagram", h1))
        if diagram_caption:
            story.append(Paragraph(escape(diagram_caption), body))
        if diagram_image_path and diagram_image_path.exists():
            story.append(Image(str(diagram_image_path), width=6.2 * inch, height=3.6 * inch, kind="proportional"))
        else:
            story.append(Paragraph(escape(diagram_source).replace("\n", "<br/>"), code_style))

    story.append(PageBreak())
    story.append(Paragraph("Approved Improvement Plan", h1))
    approved_items = approval.get("approved_items", []) if approval else []
    if approved_items:
        for item in approved_items:
            story.append(Paragraph(escape(str(item.get("title", "Improvement item"))), h2))
            story.append(Paragraph(escape(str(item.get("action", ""))), body))
    else:
        story.append(Paragraph("No approved improvement items were recorded for this export.", body))

    story.append(PageBreak())
    story.append(Paragraph(escape(certificate["title"]), h1))
    story.append(Paragraph(escape(certificate["statement"]), body))
    story.append(Paragraph(escape(f"Generated: {certificate['generated_at']}"), body))
    score_rows = [["Metric", "Before", "After"]]
    keys = sorted(set(certificate["before"].keys()) | set(certificate["after"].keys()))
    for key in keys:
        score_rows.append([
            key.replace("_", " ").title(),
            str(certificate["before"].get(key, "")),
            str(certificate["after"].get(key, "")),
        ])
    score_table = Table(score_rows, colWidths=[2.7 * inch, 1.4 * inch, 1.4 * inch])
    score_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), accent),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d1d5db")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(score_table)

    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=document_title,
    )
    pdf.build(story)
    return ExportArtifact("pdf", output_path.name, output_path, output_path.stat().st_size)


def create_certificate_markdown(
    *,
    output_path: Path,
    document_title: str,
    certificate: dict[str, Any],
    approval: dict[str, Any] | None,
) -> ExportArtifact:
    lines = [
        f"# {certificate['title']}",
        "",
        f"Document: {document_title}",
        f"Generated: {certificate['generated_at']}",
        f"Target format: {certificate['target_format']}",
        f"Chapters finalized: {certificate['chapter_count']}",
        "",
        certificate["statement"],
        "",
        "## Approval Scope",
    ]
    approved_items = approval.get("approved_items", []) if approval else []
    if approved_items:
        for item in approved_items:
            lines.append(f"- {item.get('title', 'Improvement item')}: {item.get('action', '')}")
    else:
        lines.append("- No approved improvement items were recorded.")
    lines.extend(["", "## Before / After Scores"])
    keys = sorted(set(certificate["before"].keys()) | set(certificate["after"].keys()))
    for key in keys:
        lines.append(f"- {key}: {certificate['before'].get(key, '')} -> {certificate['after'].get(key, '')}")
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return ExportArtifact("md", output_path.name, output_path, output_path.stat().st_size)
